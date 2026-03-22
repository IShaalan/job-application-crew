#!/usr/bin/env python3
"""
Convert a markdown resume or cover letter to a formatted DOCX file.

Uses python-docx to create a Word document with basic professional formatting:
- Name in 16pt bold
- Title in 12pt
- Contact line in 10pt
- Section headers in 12pt bold
- Bullets in 10.5pt
- 0.7 inch margins

Usage:
    python scripts/generate_docx.py --content final/resume.md --type resume --company Acme --role "Senior PM"
    python scripts/generate_docx.py --content final/cover-letter.md --type cover-letter --company Acme --role "Senior PM"
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

try:
    import yaml
except ImportError:
    yaml = None


PROJECT_ROOT = Path(__file__).parent.parent


def load_candidate_name(config_path=None):
    """Load candidate name from resume-config.yaml.

    Returns the candidate's name for the output filename, or 'Output' as default.
    """
    if yaml is None:
        return "Output"

    if config_path:
        paths = [Path(config_path)]
    else:
        paths = [
            PROJECT_ROOT / "candidate" / "resume-config.yaml",
            PROJECT_ROOT / "candidate" / "resume-config.yml",
        ]

    for path in paths:
        if path.exists():
            try:
                with open(path) as f:
                    config = yaml.safe_load(f) or {}
                # Try common config structures
                contact = config.get("contact", {})
                name = contact.get("name") or contact.get("full_name")
                if name:
                    return name
                # Top-level name field
                name = config.get("name") or config.get("candidate_name")
                if name:
                    return name
            except Exception:
                pass

    return "Output"


def sanitize_filename(text):
    """Remove spaces and special characters from a string for use in filenames."""
    return re.sub(r'[^a-zA-Z0-9]', '', text)


def build_output_filename(name, doc_type, company, role):
    """Build output filename: {Name}-{DocType}-{Company}-{Role}.docx"""
    name_part = sanitize_filename(name)
    type_part = "Resume" if doc_type == "resume" else "CoverLetter"
    company_part = sanitize_filename(company)
    role_part = sanitize_filename(role)
    return f"{name_part}-{type_part}-{company_part}-{role_part}.docx"


def set_margins(document, margin_inches=0.7):
    """Set all page margins to the specified value."""
    for section in document.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)


def add_formatted_text(paragraph, text, size_pt, bold=False):
    """Add a run of text to a paragraph with specified formatting."""
    run = paragraph.add_run(text)
    run.font.size = Pt(size_pt)
    run.bold = bold
    return run


def strip_markdown_links(text):
    """Convert [text](url) markdown links to just the display text."""
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def parse_inline_bold(text):
    """Parse text with **bold** markers into segments.

    Returns a list of (text, is_bold) tuples.
    """
    segments = []
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            segments.append((part[2:-2], True))
        else:
            segments.append((part, False))
    return segments


def add_paragraph_with_inline_bold(document, text, base_size_pt, base_bold=False):
    """Add a paragraph handling **bold** inline markers."""
    paragraph = document.add_paragraph()
    segments = parse_inline_bold(text)
    for segment_text, is_bold in segments:
        run = paragraph.add_run(segment_text)
        run.font.size = Pt(base_size_pt)
        run.bold = base_bold or is_bold
    return paragraph


def parse_markdown(content):
    """Parse markdown content into structured elements.

    Returns a list of dicts with keys:
        - type: 'h1', 'h2', 'title_bold', 'contact', 'bullet', 'text', 'blank'
        - text: the content
    """
    elements = []
    lines = content.strip().split('\n')
    found_h1 = False
    found_title = False
    found_contact = False

    for i, line in enumerate(lines):
        stripped = line.strip()

        # Blank line
        if not stripped:
            elements.append({'type': 'blank', 'text': ''})
            continue

        # H1 header (# Name)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            elements.append({'type': 'h1', 'text': stripped[2:].strip()})
            found_h1 = True
            continue

        # H2 header (## Section)
        if stripped.startswith('## '):
            elements.append({'type': 'h2', 'text': stripped[3:].strip()})
            continue

        # Right after H1: look for bold title line and contact line
        if found_h1 and not found_title:
            # Bold title line: **Title** or **Title Text**
            bold_match = re.match(r'^\*\*(.+)\*\*$', stripped)
            if bold_match:
                elements.append({'type': 'title_bold', 'text': bold_match.group(1)})
                found_title = True
                continue

        if found_h1 and not found_contact:
            # Contact line: contains email pattern or pipe separators typical of contact lines
            if ('|' in stripped and '@' in stripped) or ('@' in stripped and '+' in stripped):
                clean = strip_markdown_links(stripped)
                elements.append({'type': 'contact', 'text': clean})
                found_contact = True
                continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            bullet_text = stripped[2:].strip()
            elements.append({'type': 'bullet', 'text': bullet_text})
            continue

        # Bold line (e.g., experience entry: **Title | Company | Location | Dates**)
        bold_match = re.match(r'^\*\*(.+)\*\*$', stripped)
        if bold_match:
            elements.append({'type': 'bold_line', 'text': bold_match.group(1)})
            continue

        # Regular text
        elements.append({'type': 'text', 'text': stripped})

    return elements


def generate_docx(elements, output_path):
    """Generate a DOCX file from parsed markdown elements."""
    document = Document()
    set_margins(document)

    # Set default font
    style = document.styles['Normal']
    style.font.size = Pt(10.5)

    for elem in elements:
        etype = elem['type']
        text = elem['text']

        if etype == 'blank':
            # Add minimal spacing — skip adding empty paragraphs to keep it tight
            continue

        elif etype == 'h1':
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(p, text, 16, bold=True)
            p.paragraph_format.space_after = Pt(2)

        elif etype == 'title_bold':
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(p, text, 12, bold=False)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)

        elif etype == 'contact':
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(p, text, 10, bold=False)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)

        elif etype == 'h2':
            p = document.add_paragraph()
            add_formatted_text(p, text, 12, bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)

        elif etype == 'bullet':
            p = document.add_paragraph(style='List Bullet')
            # Clear default text and add with formatting + inline bold support
            segments = parse_inline_bold(text)
            for seg_text, is_bold in segments:
                run = p.add_run(seg_text)
                run.font.size = Pt(10.5)
                run.bold = is_bold
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

        elif etype == 'bold_line':
            p = document.add_paragraph()
            add_formatted_text(p, text, 10.5, bold=True)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)

        elif etype == 'text':
            p = add_paragraph_with_inline_bold(document, text, 10.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

    document.save(str(output_path))
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert markdown resume/cover letter to DOCX'
    )
    parser.add_argument(
        '--content', required=True,
        help='Path to markdown file'
    )
    parser.add_argument(
        '--type', required=True, choices=['resume', 'cover-letter'],
        help='Document type'
    )
    parser.add_argument(
        '--company', required=True,
        help='Company name'
    )
    parser.add_argument(
        '--role', required=True,
        help='Role name'
    )
    parser.add_argument(
        '--output',
        help='Output path (default: same directory as input)'
    )
    parser.add_argument(
        '--config',
        help='Path to resume-config.yaml'
    )

    args = parser.parse_args()

    # Validate input file
    content_path = Path(args.content)
    if not content_path.exists():
        print(f"Error: Input file not found: {args.content}")
        sys.exit(1)

    # Load candidate name
    candidate_name = load_candidate_name(args.config)

    # Build output path
    filename = build_output_filename(candidate_name, args.type, args.company, args.role)

    if args.output:
        output_path = Path(args.output)
        # If output is a directory, put the file inside it
        if output_path.is_dir():
            output_path = output_path / filename
    else:
        output_path = content_path.parent / filename

    # Read and parse markdown
    content = content_path.read_text(encoding='utf-8')
    elements = parse_markdown(content)

    if not elements:
        print("Error: No content parsed from the markdown file.")
        sys.exit(1)

    # Generate DOCX
    result_path = generate_docx(elements, output_path)
    print(f"Generated: {result_path}")


if __name__ == '__main__':
    main()
